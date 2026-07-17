import logging
import io
import time
from typing import Tuple, List, Dict, Any, AsyncIterator
from docx import Document

from core.exceptions import FileProcessingException
from .base_extractor import ITextExtractor

logger = logging.getLogger(__name__)

class WordTextExtractor(ITextExtractor):
    """Extract text from Word document."""

    async def extract_legacy(self, content: bytes) -> Tuple[List[str], float, Dict[str, Any]]:
        if not content:
            raise FileProcessingException("word", "Empty DOCX provided")
        
        try:
            start_time = time.time()
            doc = Document(io.BytesIO(content))
            
            all_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    all_text.append(para.text)
            
            confidence = 1.0 if all_text else 0.0
            processing_time = time.time() - start_time
            
            log_entry = {
                "source": "word_native",
                "success": True,
                "paragraphs": len(doc.paragraphs),
                "text_length": sum(len(t) for t in all_text),
                "confidence": confidence,
                "processing_time": processing_time
            }
            
            return all_text, confidence, log_entry
        except Exception as e:
            logger.error(f"Word extraction failed: {e}")
            raise FileProcessingException("word", f"Extraction failed: {e}")

    async def extract_stream(self, content: bytes) -> AsyncIterator[Tuple[str, float, Dict[str, Any]]]:
        if not content:
            raise FileProcessingException("word", "Empty DOCX provided")
            
        try:
            doc = Document(io.BytesIO(content))
            total_paras = len(doc.paragraphs)
            for idx, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    logger.info(f"Extracted native text from Word paragraph {idx + 1}: {text[:100]}...")
                    yield text, 1.0, {
                        "source": "word_native_paragraph",
                        "success": True,
                        "paragraph_index": idx + 1,
                        "total_paragraphs": total_paras,
                        "text_length": len(text),
                    }
        except Exception as e:
            logger.error(f"Word extraction failed: {e}")
            raise FileProcessingException("word", f"Extraction failed: {e}")
